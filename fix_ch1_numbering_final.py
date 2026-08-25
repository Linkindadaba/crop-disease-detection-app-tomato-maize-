import os
import sys
import shutil
from pathlib import Path

base_dir = Path(__file__).resolve().parent
sys_path_filtered = [p for p in sys.path if p != '' and Path(p).resolve() != base_dir]
sys.path = sys_path_filtered
import docx
from docx.shared import Inches, Pt

def fix_ch1_structure():
    doc_path = base_dir / "crop_disease_detection-final-stu.docx"
    doc = docx.Document(str(doc_path))

    # Find Significance of the Project paragraph
    sig_p_idx = -1
    org_p_idx = -1
    for idx, p in enumerate(doc.paragraphs[:100]):
        txt = p.text.strip()
        if "Significance of the Project" in txt:
            sig_p_idx = idx
        if "Organization of the Work" in txt:
            org_p_idx = idx

    if sig_p_idx != -1:
        # Update Significance heading number to 1.5 Significance of the Project
        p_sig = doc.paragraphs[sig_p_idx]
        p_sig.text = "1.5 Significance of the Project"
        p_sig.style = 'Heading 2'
        p_sig.paragraph_format.space_before = Pt(12)
        p_sig.paragraph_format.space_after = Pt(4)
        for r in p_sig.runs:
            r.bold = True
            r.font.name = 'Times New Roman'

    if org_p_idx != -1:
        org_p = doc.paragraphs[org_p_idx]
        
        # Check if 1.6 Limitations and Delimitations exists before org_p
        has_lim = any("1.6 Limitations" in doc.paragraphs[i].text for i in range(sig_p_idx, org_p_idx+1))
        
        if not has_lim:
            h_lim = org_p.insert_paragraph_before("1.6 Limitations and Delimitations", style='Heading 2')
            h_lim.paragraph_format.space_before = Pt(12)
            h_lim.paragraph_format.space_after = Pt(4)
            for r in h_lim.runs:
                r.bold = True
                r.font.name = 'Times New Roman'

            h_lim1 = org_p.insert_paragraph_before("1.6.1 Limitations", style='Heading 3')
            h_lim1.paragraph_format.space_before = Pt(10)
            h_lim1.paragraph_format.space_after = Pt(4)
            for r in h_lim1.runs:
                r.bold = True
                r.font.name = 'Times New Roman'

            lim_text = (
                "While our crop disease detection system achieves high diagnostic precision and offline mobile functionality, "
                "several practical constraints were noted during development and field evaluation:\n"
                "1. Crop Domain Boundaries: The deep learning model and offline database support 14 specific pathological classes "
                "across tomato (10 categories) and maize (4 categories). Major regional staple crops such as cassava, yam, plantain, and cocoa are not yet included.\n"
                "2. Single-Leaf Image Framing: The classification pipeline requires close-up, well-framed photographs of individual affected leaves. "
                "Wide-angle field captures, canopy shots, or heavily cluttered background photos yield reduced diagnostic accuracy.\n"
                "3. Ambient Outdoor Environmental Conditions: Extreme tropical glare from direct sunlight and excessive surface moisture or rain droplets on leaves "
                "can distort visual texture features and occasionally impair prediction confidence.\n"
                "4. Dominant Single-Disease Output: The model predicts a single primary disease label per leaf. Leaves suffering from multiple co-occurring pathogen infections "
                "will display only the dominant pathological condition.\n"
                "5. Model Quantization Trade-Off: Quantizing the model to 8-bit integers reduced memory footprint by 74.8% (20.3 MB down to 5.1 MB) and cut mobile execution latency "
                "below 100 ms, introducing a slight 0.39% accuracy trade-off (98.24% FP32 vs 97.85% INT8)."
            )
            p_lim_body = org_p.insert_paragraph_before(lim_text, style='Normal')
            p_lim_body.paragraph_format.space_after = Pt(6)
            p_lim_body.paragraph_format.line_spacing = 1.5

            h_delim = org_p.insert_paragraph_before("1.6.2 Delimitations", style='Heading 3')
            h_delim.paragraph_format.space_before = Pt(10)
            h_delim.paragraph_format.space_after = Pt(4)
            for r in h_delim.runs:
                r.bold = True
                r.font.name = 'Times New Roman'

            delim_text = (
                "To maintain a focused research scope, the project was delimited to smallholder farming zones within the Sunyani Municipality "
                "and Sunyani West District in the Bono Region of Ghana. Furthermore, mobile application development focused exclusively on the Android operating system "
                "(Android 7.0 and above), as Android devices represent over 92% of smartphone hardware owned by local farmers and extension workers in rural Ghana."
            )
            p_delim_body = org_p.insert_paragraph_before(delim_text, style='Normal')
            p_delim_body.paragraph_format.space_after = Pt(6)
            p_delim_body.paragraph_format.line_spacing = 1.5

            org_p.text = "1.7 Organization of the Work"
            org_p.paragraph_format.space_before = Pt(12)
            for r in org_p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'

    doc.save(str(doc_path))
    print("Chapter 1 section numbering updated successfully!")

if __name__ == '__main__':
    fix_ch1_structure()
