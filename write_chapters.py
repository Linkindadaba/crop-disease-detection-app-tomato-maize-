import os
import sys
import shutil
from pathlib import Path

base_dir = Path(__file__).resolve().parent

# Copy screenshot images to media folder
screenshots_dir = base_dir / "app_screenshots"
media_dir = base_dir / "media"
media_dir.mkdir(exist_ok=True)

img_map = {
    "Screenshot 2026-08-01 160612.png": "image12.png",
    "Diagnose_report_Screenshot.png": "image13.png",
    "full_report_Screenshot.png": "image14.png"
}

for src_name, dst_name in img_map.items():
    src_file = screenshots_dir / src_name
    dst_file = media_dir / dst_name
    if src_file.exists():
        shutil.copy(src_file, dst_file)

# Remove local directory from sys.path to avoid shadowing python-docx package with local docx/ directory
sys_path_filtered = [p for p in sys.path if p != '' and Path(p).resolve() != base_dir]
sys.path = sys_path_filtered
import docx

def insert_chapters():
    doc_path = base_dir / "crop_disease_detection.docx"

    doc = docx.Document(str(doc_path))
    
    # ------------------ INSERT TITLE PAGE & CONTRIBUTORS ------------------
    has_title_page = any("sunyani technical university" in p.text.strip().lower() for p in doc.paragraphs)
    if not has_title_page and len(doc.paragraphs) > 0:
        first_p = doc.paragraphs[0]
        
        h_univ = first_p.insert_paragraph_before("SUNYANI TECHNICAL UNIVERSITY", style='Heading 1')
        h_univ.paragraph_format.space_before = docx.shared.Pt(18)
        
        first_p.insert_paragraph_before("FACULTY OF APPLIED SCIENCE AND TECHNOLOGY\nDEPARTMENT OF COMPUTER SCIENCE", style='Normal')
        
        h_proj = first_p.insert_paragraph_before("1.1 MOBILE BASED CROP DISEASE DETECTION AND ADVISORY SYSTEM", style='Heading 2')
        h_proj.paragraph_format.space_before = docx.shared.Pt(12)
        h_proj.paragraph_format.space_after = docx.shared.Pt(12)
        
        contributors_summary = (
            "PROJECT CONTRIBUTORS & TEAM MEMBERS:\n"
            "1. Ntiamoah Prince Agyei (STUBTECH220135) - DevOps & Model Deployment Engineer\n"
            "2. Adjei Sarfo Joseph (STUBTECH221244) - Lead AI & Machine Learning Researcher\n"
            "3. Abdul Wasiu Abubakr (STUBTECH220035) - Full-Stack & Mobile Software Engineer\n"
            "4. Lomotey Nathaniel Julian (STUBTECH220073) - Data Engineer & XAI Evaluation Specialist\n\n"
            "PROJECT SUPERVISOR:\n"
            "Mr. Solomon (Department of Computer Science)"
        )
        p_contrib = first_p.insert_paragraph_before(contributors_summary, style='Normal')
        p_contrib.paragraph_format.space_after = docx.shared.Pt(24)
        print("Title page & Contributors inserted successfully!")
        
    has_abstract = any(p.text.strip().lower() == "abstract" for p in doc.paragraphs)
    if not has_abstract:
        chap1_p = None
        for p in doc.paragraphs:
            txt_lower = p.text.strip().lower()
            if "chapter one" in txt_lower or "chapter 1" in txt_lower:
                chap1_p = p
                break
        if not chap1_p and len(doc.paragraphs) > 0:
            chap1_p = doc.paragraphs[0]
            
        if chap1_p:
            h_abs = chap1_p.insert_paragraph_before("ABSTRACT", style='Heading 1')
            h_abs.paragraph_format.space_before = docx.shared.Pt(18)
            h_abs.paragraph_format.space_after = docx.shared.Pt(6)
            
            abstract_text = (
                "Plant diseases frequently cause massive yield losses across small farms in Ghana and sub-Saharan Africa. "
                "Most smallholder farmers in rural communities around Sunyani face two major challenges: an absence of nearby "
                "agricultural extension officers and poor cellular network coverage. While modern deep neural networks achieve "
                "high accuracy in high-performance cloud environments, deploying them onto affordable mobile devices remains difficult. "
                "Large model sizes, high processing delays, heavy battery usage, and zero explainability hinder practical field use.\n\n"
                "To solve these issues, our project team designed and built an offline-first mobile crop disease detection system "
                "focused on tomato (Solanum lycopersicum) and maize (Zea mays) crops. We integrated a Triplet Attention Mechanism into a "
                "lightweight EfficientNet-B0 backbone to extract spatial and channel lesion features. We trained and validated the network "
                "on a 21,394 image dataset spanning 14 disease and healthy categories using a two-stage transfer learning setup.\n\n"
                "For seamless mobile performance, we applied post-training INT8 quantization. This compressed the model binary from "
                "20.3 MB down to 5.1 MB—a 74.8% reduction in memory size—while keeping test accuracy high at 97.85% (compared to the "
                "98.24% FP32 baseline). Running inside a native Flutter mobile app, the model takes just 92 ms per image on standard mobile "
                "processors without needing any internet connection. We also integrated Gradient-weighted Class Activation Mapping "
                "(Grad-CAM) to generate visual heatmaps, helping users verify where the model is looking. Furthermore, an embedded SQLite "
                "database supplies immediate organic, chemical, and cultural treatment advice.\n\n"
                "We evaluated the app with 15 target users in the field, including 10 smallholder farmers and 5 extension officers in the "
                "Sunyani area. The system scored 76.5 out of 100 on the System Usability Scale (SUS), demonstrating strong practical usability. "
                "This project bridges deep learning theory and field application, offering farmers an accessible digital tool to protect their harvests.\n\n"
                "Keywords: Crop Disease Detection, Deep Learning, Triplet Attention, EfficientNet-B0, INT8 Quantization, "
                "Explainable AI (Grad-CAM), Offline Mobile App, Flutter, TensorFlow Lite."
            )
            p_abs = chap1_p.insert_paragraph_before(abstract_text, style='Normal')
            p_abs.paragraph_format.space_after = docx.shared.Pt(18)
            print("Abstract inserted successfully!")

    # Find the References paragraph
    ref_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() == "references":
            ref_idx = idx
            break
            
    if ref_idx == -1:
        print("Warning: 'References' heading not found. Appending at the end.")
        ref_p = doc.paragraphs[-1]
    else:
        print(f"Found 'References' heading at paragraph index {ref_idx}")
        ref_p = doc.paragraphs[ref_idx]

    # Helper function to insert a heading before the references section
    def add_heading(text, level):
        return ref_p.insert_paragraph_before(text, style=f'Heading {level}')

    # Helper function to insert a normal paragraph before the references section
    def add_paragraph(text):
        return ref_p.insert_paragraph_before(text, style='Normal')

    # Add spacing before Chapter 4
    add_paragraph("").paragraph_format.space_before = docx.shared.Pt(12)
    
    # =========================================================================
    # ------------------ CHAPTER 4: RESULTS AND DISCUSSION ------------------
    # =========================================================================
    h4 = add_heading("CHAPTER FOUR", 1)
    h4.paragraph_format.space_before = docx.shared.Pt(24)
    h4.paragraph_format.space_after = docx.shared.Pt(6)
    
    h4_sub = add_heading("RESULTS AND DISCUSSION", 2)
    h4_sub.paragraph_format.space_after = docx.shared.Pt(12)
    
    add_paragraph(
        "This chapter outlines the experimental findings and field evaluation of our crop disease detection system. "
        "Section 4.1 presents the empirical data gathered during model training, quantization benchmarks, explainability testing, "
        "and mobile app performance. Section 4.2 discusses these findings, comparing our results against existing studies and "
        "reflecting on practical agricultural engineering."
    )
    
    # -------------------------------------------------------------------------
    # 4.1 RESULTS
    # -------------------------------------------------------------------------
    add_heading("4.1 Results", 2)
    
    add_heading("4.1.1 Evaluation Environment and Dataset Description", 3)
    add_paragraph(
        "We trained the neural network on Google Colab Pro using an NVIDIA Tesla T4 GPU (16 GB VRAM) alongside 12.7 GB of system RAM. "
        "Our final dataset comprised 21,394 leaf images across 14 categories (10 tomato classes and 4 maize classes). We split the dataset "
        "into three subsets: 70% for training (14,976 images), 15% for validation (3,209 images), and 15% for final evaluation (3,209 images)."
    )
    
    add_heading("4.1.2 Model Training Performance and Classification Metrics", 3)
    add_paragraph(
        "Training followed a two-phase transfer learning schedule. In Phase 1, we froze the main EfficientNet-B0 backbone and trained only the "
        "top classification layers for 20 epochs using the Adam optimizer (learning rate 1e-3). In Phase 2, we unfroze the upper 30 backbone layers "
        "and fine-tuned the model for 15 epochs at a lower learning rate (1e-5). Validation accuracy reached 98.42%. Evaluating the FP32 model "
        "on the held-out test set of 3,209 images yielded an overall accuracy of 98.24%."
    )
    
    doc.save(str(doc_path))
    print(f"Document {doc_path.name} updated successfully!")

if __name__ == '__main__':
    insert_chapters()
