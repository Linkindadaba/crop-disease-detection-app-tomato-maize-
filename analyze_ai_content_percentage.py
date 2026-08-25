import sys
import re
import math
from pathlib import Path

base_dir = Path(__file__).resolve().parent
sys_path_filtered = [p for p in sys.path if p != '' and Path(p).resolve() != base_dir]
sys.path = sys_path_filtered
import docx

# List of typical AI transition markers & buzzwords targeted by AI detectors (Turnitin, ZeroGPT, CopyLeaks)
AI_BUZZWORDS = [
    "delve", "delves", "delving",
    "testament", "pivotal", "paradigm shift",
    "cutting-edge", "robust", "comprehensive",
    "furthermore", "moreover", "additionally",
    "consequently", "it is worth noting", "it should be noted",
    "underscores", "tapestry", "beacon", "seamless",
    "realm", "landscape", "paving the way", "vital role",
    "in conclusion", "to summarize", "transformative",
    "game-changer", "empower", "harnessing", "unraveling",
    "interplay", "multifaceted", "holistic", "crucial role"
]

def analyze_document():
    doc_path = base_dir / "crop_disease_detection.docx"
    if not doc_path.exists():
        print("Error: Document crop_disease_detection.docx not found.")
        return

    doc = docx.Document(str(doc_path))
    
    total_paragraphs = len(doc.paragraphs)
    total_words = 0
    sentences = []
    word_freq = {}
    ai_phrase_matches = []
    high_ai_paragraphs = []

    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        words = re.findall(r'\b[a-zA-Z0-9-]+\b', text.lower())
        if not words:
            continue
            
        total_words += len(words)
        for w in words:
            word_freq[w] = word_freq.get(w, 0) + 1
            
        p_sentences = [s.strip() for s in re.split(r'[.!?]+', text) if len(s.strip().split()) > 3]
        sentences.extend(p_sentences)

        # Check for AI phrases in paragraph
        para_ai_hits = []
        for word_or_phrase in AI_BUZZWORDS:
            pattern = r'\b' + re.escape(word_or_phrase) + r'\b'
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                para_ai_hits.extend(matches)

        if para_ai_hits:
            ai_phrase_matches.extend(para_ai_hits)
            if len(para_ai_hits) >= 2:
                high_ai_paragraphs.append((idx, text[:80] + "...", len(para_ai_hits)))

    # Metrics calculation
    # 1. Type-Token Ratio (Vocabulary Diversity)
    unique_words = len(word_freq)
    ttr = (unique_words / total_words) * 100 if total_words > 0 else 0

    # 2. Sentence Length Burstiness (Standard Deviation of sentence length)
    sentence_lengths = [len(s.split()) for s in sentences]
    avg_sentence_len = sum(sentence_lengths) / len(sentence_lengths) if sentence_lengths else 0
    variance = sum((l - avg_sentence_len) ** 2 for l in sentence_lengths) / len(sentence_lengths) if sentence_lengths else 0
    std_dev_sentence_len = math.sqrt(variance)

    # 3. AI Buzzword Density per 1000 words
    buzzword_density = (len(ai_phrase_matches) / total_words) * 1000 if total_words > 0 else 0

    # 4. Estimated AI Probability Score Formula
    # Standard AI text has low sentence length variance (burstiness < 6.0), high buzzword density (> 10 per 1k words), low TTR (< 18%)
    base_score = 15.0 # baseline human noise
    
    # Sentence length uniformity penalty (Low burstiness = AI signature)
    if std_dev_sentence_len < 6.0:
        base_score += 25.0
    elif std_dev_sentence_len < 9.0:
        base_score += 15.0
    elif std_dev_sentence_len < 12.0:
        base_score += 5.0
        
    # Buzzword density score
    if buzzword_density > 15:
        base_score += 35.0
    elif buzzword_density > 10:
        base_score += 25.0
    elif buzzword_density > 5:
        base_score += 15.0
    elif buzzword_density > 2:
        base_score += 5.0

    # Vocabulary diversity adjustment
    if ttr < 15.0:
        base_score += 15.0
    elif ttr < 20.0:
        base_score += 5.0

    estimated_ai_percentage = min(99.0, max(5.0, round(base_score, 1)))

    print("==================================================")
    print("        AI CONTENT & STYLISTIC ANALYSIS           ")
    print("==================================================")
    print(f"Total Words Analyzed: {total_words}")
    print(f"Total Sentences Analyzed: {len(sentences)}")
    print(f"Average Sentence Length: {avg_sentence_len:.1f} words")
    print(f"Sentence Length Burstiness (Std Dev): {std_dev_sentence_len:.2f}")
    print(f"Vocabulary Diversity (Type-Token Ratio): {ttr:.2f}%")
    print(f"Total AI Detector Marker Matches: {len(ai_phrase_matches)}")
    print(f"AI Marker Density: {buzzword_density:.2f} matches per 1,000 words")
    print("--------------------------------------------------")
    print(f"ESTIMATED AI CONTENT DETECTABILITY SCORE: {estimated_ai_percentage}%")
    print("--------------------------------------------------")
    
    print("\nMost Frequent AI Buzzwords / Detector Triggers Found:")
    from collections import Counter
    counts = Counter([m.lower() for m in ai_phrase_matches])
    for phrase, cnt in counts.most_common(12):
        print(f"  - '{phrase}': {cnt} occurrences")

    if high_ai_paragraphs:
        print(f"\nParagraphs with Highest Detector Triggers ({len(high_ai_paragraphs)} paragraphs):")
        for p_idx, snippet, hits in high_ai_paragraphs[:8]:
            print(f"  [Para #{p_idx}] Hits: {hits} | '{snippet}'")

if __name__ == '__main__':
    analyze_document()
