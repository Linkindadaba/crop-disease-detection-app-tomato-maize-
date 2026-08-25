import os
import sys
import shutil
from pathlib import Path

base_dir = Path(__file__).resolve().parent

# Ensure images exist in media folder
screenshots_dir = base_dir / "app_screenshots"
media_dir = base_dir / "media"
media_dir.mkdir(exist_ok=True)

img_map = {
    "Screenshot 2026-08-01 160612.png": "image12.png",
    "Diagnose_report_Screenshot.png": "image13.png",
    "full_report_Screenshot.png": "image14.png"
}

for src_name, dst_name in img_map.items():
    src_file = screenshots_dir / src_name
    dst_file = media_dir / dst_name
    if src_file.exists():
        shutil.copy(src_file, dst_file)

# Filter sys.path to avoid local docx folder shadowing python-docx package
sys_path_filtered = [p for p in sys.path if p != '' and Path(p).resolve() != base_dir]
sys.path = sys_path_filtered
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_document_margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.5)
        s.right_margin = Inches(1.0)

def apply_scientific_research_track():
    doc_path = base_dir / "crop_disease_detection.docx"
    doc = docx.Document(str(doc_path))
    format_document_margins(doc)

    # Clean up chapter titles
    for i in range(len(doc.paragraphs)):
        txt_upper = doc.paragraphs[i].text.strip().upper()
        if txt_upper in ["CHAPTER THREE", "CHAPTER 3"]:
            doc.paragraphs[i].text = "CHAPTER THREE"
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].text = "METHODOLOGY"
        elif txt_upper in ["CHAPTER FOUR", "CHAPTER 4"]:
            doc.paragraphs[i].text = "CHAPTER FOUR"
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].text = "RESULTS AND DISCUSSION"

    # Relocate misplaced section 3.12-3.14 to before Chapter Four if needed
    ch4_idx = -1
    sec312_idx = -1
    appA_idx = -1

    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().upper()
        if txt == "CHAPTER FOUR" and ch4_idx == -1:
            ch4_idx = idx
        if "3.12" in txt and "LOCAL OPERATIONAL CONTEXT" in txt and sec312_idx == -1:
            sec312_idx = idx
        if "APPENDIX A:" in txt and appA_idx == -1:
            appA_idx = idx

    if ch4_idx != -1 and sec312_idx != -1 and appA_idx != -1 and sec312_idx > ch4_idx:
        print(f"Relocating 3.12-3.14 (paragraphs {sec312_idx} to {appA_idx}) to before Chapter Four...")
        ch4_p = doc.paragraphs[ch4_idx]
        paras_to_move = doc.paragraphs[sec312_idx:appA_idx]
        elems = [p._element for p in paras_to_move]
        for elem in elems:
            ch4_p._element.addprevious(elem)

    doc.save(str(doc_path))
    print(f"Applied Scientific Research Track titles to document! Paragraphs: {len(doc.paragraphs)}")

if __name__ == '__main__':
    apply_scientific_research_track()
