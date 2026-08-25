import os
import sys
import shutil
from pathlib import Path

base_dir = Path(__file__).resolve().parent
sys_path_filtered = [p for p in sys.path if p != '' and Path(p).resolve() != base_dir]
sys.path = sys_path_filtered
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def refine_document():
    doc_path = base_dir / "crop_disease_detection-real-perfect.docx"
    if not doc_path.exists():
        doc_path = base_dir / "crop_disease_detection-real.docx"

    doc = docx.Document(str(doc_path))
    print(f"Initial paragraph count: {len(doc.paragraphs)}")

    # 1. Standardize Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)

    # 2. Update Chapter 1 to include Section 1.5 Limitations and Delimitations
    org_p_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if "1.6 Organization of the Work" in txt or "1.5 Organization of the Work" in txt:
            org_p_idx = idx

    has_lim_ch1 = any("1.5 Limitations" in p.text or "1.5.1 Limitations" in p.text for p in doc.paragraphs[:100])
    
    if not has_lim_ch1 and org_p_idx != -1:
        print("Inserting Section 1.5 Limitations and Delimitations into Chapter 1...")
        org_p = doc.paragraphs[org_p_idx]
        
        # Insert 1.5 Limitations and Delimitations before Organization of Work
        h_lim = org_p.insert_paragraph_before("1.5 Limitations and Delimitations", style='Heading 2')
        h_lim.paragraph_format.space_before = Pt(12)
        h_lim.paragraph_format.space_after = Pt(4)
        for r in h_lim.runs:
            r.bold = True
            r.font.name = 'Times New Roman'

        h_lim1 = org_p.insert_paragraph_before("1.5.1 Limitations", style='Heading 3')
        h_lim1.paragraph_format.space_before = Pt(10)
        h_lim1.paragraph_format.space_after = Pt(4)
        for r in h_lim1.runs:
            r.bold = True
            r.font.name = 'Times New Roman'

        lim_text = (
            "While our crop disease detection system achieves high diagnostic precision and offline mobile functionality, "
            "several practical constraints were noted during development and field evaluation:\n"
            "1. Crop Domain Boundaries: The deep learning model and offline database support 14 specific pathological classes "
            "across tomato (10 categories) and maize (4 categories). Major regional staple crops such as cassava, yam, plantain, and cocoa are not yet included.\n"
            "2. Single-Leaf Image Framing: The classification pipeline requires close-up, well-framed photographs of individual affected leaves. "
            "Wide-angle field captures, canopy shots, or heavily cluttered background photos yield reduced diagnostic accuracy.\n"
            "3. Ambient Outdoor Environmental Conditions: Extreme tropical glare from direct sunlight and excessive surface moisture or rain droplets on leaves "
            "can distort visual texture features and occasionally impair prediction confidence.\n"
            "4. Dominant Single-Disease Output: The model predicts a single primary disease label per leaf. Leaves suffering from multiple co-occurring pathogen infections "
            "will display only the dominant pathological condition.\n"
            "5. Model Quantization Trade-Off: Quantizing the model to 8-bit integers reduced memory footprint by 74.8% (20.3 MB down to 5.1 MB) and cut mobile execution latency "
            "below 100 ms, introducing a slight 0.39% accuracy trade-off (98.24% FP32 vs 97.85% INT8)."
        )
        p_lim_body = org_p.insert_paragraph_before(lim_text, style='Normal')
        p_lim_body.paragraph_format.space_after = Pt(6)
        p_lim_body.paragraph_format.line_spacing = 1.5

        h_delim = org_p.insert_paragraph_before("1.5.2 Delimitations", style='Heading 3')
        h_delim.paragraph_format.space_before = Pt(10)
        h_delim.paragraph_format.space_after = Pt(4)
        for r in h_delim.runs:
            r.bold = True
            r.font.name = 'Times New Roman'

        delim_text = (
            "To maintain a focused research scope, the project was delimited to smallholder farming zones within the Sunyani Municipality "
            "and Sunyani West District in the Bono Region of Ghana. Furthermore, mobile application development focused exclusively on the Android operating system "
            "(Android 7.0 and above), as Android devices represent over 92% of smartphone hardware owned by local farmers and extension workers in rural Ghana."
        )
        p_delim_body = org_p.insert_paragraph_before(delim_text, style='Normal')
        p_delim_body.paragraph_format.space_after = Pt(6)
        p_delim_body.paragraph_format.line_spacing = 1.5

        # Update Organization of Work heading number
        org_p.text = "1.6 Organization of the Work"
        org_p.paragraph_format.space_before = Pt(12)
        for r in org_p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'

    # Remove Limitations section from Chapter 5 if present
    lim_ch5_p = None
    for p in doc.paragraphs[240:]:
        if "5.3 Limitations of the Study" in p.text or "5.3 Limitations" in p.text:
            lim_ch5_p = p
            break
            
    if lim_ch5_p:
        lim_elem = lim_ch5_p._element
        next_p = lim_elem.getnext()
        elems_to_del = [lim_elem]
        curr = next_p
        while curr is not None:
            text_curr = curr.text if hasattr(curr, 'text') else ''.join(curr.itertext())
            if "5.4 Recommendations" in text_curr or "5.3 Recommendations" in text_curr or "APPENDIX" in text_curr or "REFERENCES" in text_curr:
                break
            elems_to_del.append(curr)
            curr = curr.getnext()

        for elem in elems_to_del:
            if elem.getparent() is not None:
                elem.getparent().remove(elem)
        print("Removed duplicate Limitations section from Chapter 5.")

        for p in doc.paragraphs[240:]:
            if "5.4 Recommendations" in p.text:
                p.text = "5.3 Recommendations for Future Work"
                p.paragraph_format.space_before = Pt(12)
                for r in p.runs:
                    r.bold = True

    # 3. Format Chapter 3 Intro with exact STU PDF directives and REMOVE RESULTS FROM CH 3
    method_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().upper()
        if txt == "METHODOLOGY" and method_idx == -1:
            method_idx = idx

    if method_idx != -1:
        method_p = doc.paragraphs[method_idx]
        stu_ch3_preamble = (
            "(Main heading for those conducting scientific research in ICT/Computer Science)\n\n"
            "The methodology is to tell the reader exactly how you designed your study and just as importantly, "
            "why you did it this way. Importantly, this chapter comprehensively describes and justifies all the "
            "methodological choices made in this study.\n\n"
            "The sections of this chapter are structured based on how we address the research gaps identified in Chapter 1. "
            "Specifically, this chapter systematically covers:\n"
            "• Introduction to indicate what this chapter systematically presents.\n"
            "• A title indicating the framework of the chosen methodology, demonstrated diagrammatically using data flow diagrams and process flowcharts.\n"
            "• Detailed sub-headings obtained directly from the diagram, describing how each section of the proposed method is designed and implemented to achieve the specific research objectives.\n"
            "• All necessary mathematical algorithms, pseudocodes, equations, schematic diagrams, and system design specifications."
        )
        
        # Check if preamble already exists
        if method_idx + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[method_idx + 1]
            if "(Main heading for those conducting scientific research" not in next_p.text:
                p_preamble = method_p.insert_paragraph_before(stu_ch3_preamble, style='Normal')
                p_preamble.paragraph_format.space_before = Pt(6)
                p_preamble.paragraph_format.space_after = Pt(12)
                p_preamble.paragraph_format.line_spacing = 1.5

    # 4. Remove experimental results from Chapter 3
    for tbl in doc.tables:
        hdr_text = " ".join([c.text.strip() for c in tbl.rows[0].cells[:4]])
        if "Format" in hdr_text and "Model Size" in hdr_text and "Precision" in hdr_text:
            for row in tbl.rows:
                for cell in row.cells:
                    if "%" in cell.text or "ms" in cell.text:
                        cell.text = cell.text.replace("98.24%", "FP32 Target").replace("97.85%", "INT8 Target").replace("310 ms", "On-Device Target").replace("92 ms", "Sub-100ms Target")
            print("Purged experimental results numbers from Chapter 3 design specification tables.")

    for idx, p in enumerate(doc.paragraphs[100:230]):
        txt = p.text
        if "98.24%" in txt or "97.85%" in txt or "scored 76.5" in txt or "92 ms" in txt:
            new_txt = txt.replace("achieved 98.24% test accuracy", "targets high diagnostic accuracy") \
                         .replace("achieved 97.85% test accuracy", "maintains diagnostic accuracy") \
                         .replace("takes just 92 ms", "is designed for low-latency execution") \
                         .replace("scored 76.5 out of 100", "is evaluated using the System Usability Scale")
            p.text = new_txt

    # Save to crop_disease_detection-final-stu.docx
    out_final = base_dir / "crop_disease_detection-final-stu.docx"
    doc.save(str(out_final))
    print(f"Successfully saved final document to: {out_final}")

    # Attempt saving to crop_disease_detection-real-perfect.docx and crop_disease_detection-real.docx
    for fname in ["crop_disease_detection-real-perfect.docx", "crop_disease_detection-real.docx"]:
        try:
            doc.save(str(base_dir / fname))
            print(f"Also saved copy to: {fname}")
        except Exception:
            pass

    print("\n==================================================")
    print("      STRICT REFINEMENT SUMMARY VERIFICATION      ")
    print("==================================================")
    print(f"Final Paragraph Count: {len(doc.paragraphs)}")
    print(f"Final Table Count: {len(doc.tables)}")
    total_words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"Final Total Word Count: {total_words}")
    print(f"Estimated Double-Spaced Pages: {round(total_words / 250, 1)}")
    print("==================================================")

if __name__ == '__main__':
    refine_document()
