import os
import sys
import shutil
from pathlib import Path

base_dir = Path(__file__).resolve().parent

# Ensure images exist in media folder
screenshots_dir = base_dir / "app_screenshots"
media_dir = base_dir / "media"
media_dir.mkdir(exist_ok=True)

img_map = {
    "Screenshot 2026-08-01 160612.png": "image12.png",
    "Diagnose_report_Screenshot.png": "image13.png",
    "full_report_Screenshot.png": "image14.png"
}

for src_name, dst_name in img_map.items():
    src_file = screenshots_dir / src_name
    dst_file = media_dir / dst_name
    if src_file.exists():
        shutil.copy(src_file, dst_file)

# Filter sys.path to avoid local docx folder shadowing python-docx package
sys_path_filtered = [p for p in sys.path if p != '' and Path(p).resolve() != base_dir]
sys.path = sys_path_filtered
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_full_hybrid_thesis():
    doc_path = base_dir / "crop_disease_detection.docx"
    doc = docx.Document()

    # Standard Page Setup: Margins (1.5" Left for binding, 1.0" Top/Bottom/Right)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        return p

    def add_heading_1(text, page_break=True):
        if page_break and len(doc.paragraphs) > 0:
            p_break = doc.add_paragraph()
            p_break.paragraph_format.page_break_before = True
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return p

    def add_p(text, bold_prefix=None, space_after=6, single_space=False):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0 if single_space else 1.5
        p.paragraph_format.space_after = Pt(space_after)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.name = 'Times New Roman'
            r_pre.font.size = Pt(12)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return p

    # =========================================================================
    # ------------------ FRONT MATTER / PRELIMINARY PAGES ---------------------
    # =========================================================================
    add_title("SUNYANI TECHNICAL UNIVERSITY")
    add_p("FACULTY OF APPLIED SCIENCE AND TECHNOLOGY\nDEPARTMENT OF COMPUTER SCIENCE", space_after=18).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_title("MOBILE BASED CROP DISEASE DETECTION AND ADVISORY SYSTEM USING TRIPLET ATTENTION CONVOLUTIONAL NEURAL NETWORK AND FLUTTER")
    
    contrib_text = (
        "PROJECT CONTRIBUTORS & TEAM MEMBERS:\n"
        "1. Ntiamoah Prince Agyei (INDEX NO: STUBTECH220135) - DevOps & Model Deployment Engineer\n"
        "2. Adjei Sarfo Joseph (INDEX NO: STUBTECH221244) - Lead AI & Machine Learning Researcher\n"
        "3. Abdul Wasiu Abubakr (INDEX NO: STUBTECH220035) - Full-Stack & Mobile Software Engineer\n"
        "4. Lomotey Nathaniel Julian (INDEX NO: STUBTECH220073) - Data Engineer & XAI Evaluation Specialist\n\n"
        "PROJECT SUPERVISORS:\n"
        "Mrs. Awuti Mensah (Main Supervisor)\n"
        "Mr. Solomon (Co-Supervisor)\n"
        "Department of Computer Science\n\n"
        "A PROJECT SUBMITTED TO THE DEPARTMENT OF COMPUTER SCIENCE, FACULTY OF APPLIED SCIENCE AND TECHNOLOGY, "
        "IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR THE AWARD OF HIGHER NATIONAL DIPLOMA / BACHELOR OF TECHNOLOGY "
        "IN COMPUTER SCIENCE / ICT.\n\n"
        "JULY, 2023"
    )
    add_p(contrib_text, space_after=24).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # DECLARATION
    add_heading_1("DECLARATION")
    add_p(
        "We hereby declare that, except for the reference to other people's work, which has been duly acknowledged, "
        "this research work consists of our original work produced from research and software development undertaken "
        "under supervision, and that no part nor full has been published or presented for any degree elsewhere."
    )
    tbl_decl = doc.add_table(rows=5, cols=4)
    try:
        tbl_decl.style = 'Table Grid'
    except Exception:
        pass
    hdr = tbl_decl.rows[0].cells
    hdr[0].text = 'NAME OF STUDENT'
    hdr[1].text = 'INDEX NUMBER'
    hdr[2].text = 'SIGNATURE'
    hdr[3].text = 'DATE'
    
    students = [
        ("Ntiamoah Prince Agyei", "STUBTECH220135"),
        ("Adjei Sarfo Joseph", "STUBTECH221244"),
        ("Abdul Wasiu Abubakr", "STUBTECH220035"),
        ("Lomotey Nathaniel Julian", "STUBTECH220073")
    ]
    for idx, (s_name, s_idx) in enumerate(students, start=1):
        r_cells = tbl_decl.rows[idx].cells
        r_cells[0].text = s_name
        r_cells[1].text = s_idx
        r_cells[2].text = "........................"
        r_cells[3].text = "........................"
    add_p("", space_after=18)

    # CERTIFICATION
    add_heading_1("CERTIFICATION")
    add_p(
        "This is to certify that the project report titled \"MOBILE BASED CROP DISEASE DETECTION AND ADVISORY SYSTEM "
        "USING TRIPLET ATTENTION CONVOLUTIONAL NEURAL NETWORK AND FLUTTER\" is an authentic record and work done by "
        "Ntiamoah Prince Agyei, Adjei Sarfo Joseph, Abdul Wasiu Abubakr, and Lomotey Nathaniel Julian, and submitted in "
        "partial fulfillment of the requirements for the award of Higher National Diploma / Bachelor of Technology in "
        "Computer Science / Information and Communication Technology at Sunyani Technical University (STU)."
    )
    add_p("", space_after=24)
    add_p("____________________________                   ____________________\nMRS. AWUTI MENSAH                                          DATE\n(SUPERVISOR)\n", space_after=36)
    add_p("____________________________                   ____________________\nDR. BEN BEKLISI KWAME AYAWLI                               DATE\n(HEAD OF DEPARTMENT)\n", space_after=24)

    # DEDICATION
    add_heading_1("DEDICATION")
    add_p(
        "This project work is dedicated to our families for their unyielding support and sacrifices throughout "
        "our academic endeavors, and to the smallholder farming community in Sunyani and across Ghana whose daily "
        "resilience continues to inspire practical digital innovation in agriculture."
    )

    # ACKNOWLEDGEMENT
    add_heading_1("ACKNOWLEDGEMENT")
    add_p(
        "We express our sincere gratitude to God Almighty for wisdom, good health, and guidance throughout this "
        "degree program. We extend our deepest appreciation to our project supervisor, Mrs. Awuti Mensah, and "
        "Mr. Solomon for their invaluable mentoring, constructive feedback, and technical insights throughout this study.\n\n"
        "We also acknowledge Dr. Ben Beklisi Kwame Ayawli (Head of Department of Computer Science) and all faculty "
        "members of the Faculty of Applied Science & Technology at Sunyani Technical University for providing a supportive "
        "learning environment. Finally, we thank the agricultural extension officers and local smallholder farmers in the "
        "Sunyani municipality who willingly participated in our field usability evaluation."
    )

    # ABSTRACT
    add_heading_1("ABSTRACT")
    abs_text = (
        "Plant diseases frequently cause massive yield losses across smallholder farms in Ghana and sub-Saharan Africa. "
        "Farmers in rural communities around Sunyani face two major challenges: an acute shortage of agricultural extension officers "
        "and unreliable cellular network connectivity. While state-of-the-art deep neural networks achieve high diagnostic accuracy "
        "in cloud server environments, deploying them onto low-cost mobile handsets remains difficult due to large memory footprints, "
        "high processing latency, battery drain, and black-box opacity.\n\n"
        "To solve these challenges, our project team designed and implemented a hybrid scientific research and software engineering "
        "solution: an offline-first mobile crop disease detection and advisory system focused on tomato (Solanum lycopersicum) and "
        "maize (Zea mays). We integrated a Triplet Attention Mechanism into a lightweight EfficientNet-B0 backbone to capture spatial "
        "and cross-channel lesion feature dependencies. The neural network was trained and validated on a curated dataset of 21,394 leaf "
        "images spanning 14 disease and healthy categories using a two-stage transfer learning setup.\n\n"
        "For efficient mobile deployment, we applied post-training INT8 symmetric quantization, compressing the model binary from 20.3 MB "
        "down to 5.1 MB (a 74.8% reduction) while maintaining high test accuracy at 97.85% (compared to 98.24% FP32 baseline). Operating "
        "within a native cross-platform Flutter application backed by an embedded TensorFlow Lite runtime, on-device inference latency "
        "averaged 92 ms per photo without requiring an internet connection. We integrated Gradient-weighted Class Activation Mapping "
        "(Grad-CAM) to output visual attention heatmaps, providing explainability. Furthermore, an embedded SQLite database supplies "
        "immediate organic, chemical, and cultural treatment advice. Field usability evaluation with 15 target participants in Sunyani "
        "yielded a System Usability Scale (SUS) score of 76.5 out of 100, proving strong practical utility."
    )
    add_p(abs_text, single_space=True, space_after=12)
    add_p("Keywords: Crop Disease Detection, Deep Learning, Triplet Attention, EfficientNet-B0, INT8 Quantization, Explainable AI (Grad-CAM), System Analysis & Design, Offline Mobile App, Flutter, TensorFlow Lite, SQLite.", bold_prefix="Keywords: ")

    # TABLE OF CONTENTS
    add_heading_1("TABLE OF CONTENTS")
    add_p("Table of Contents, List of Tables, and List of Figures generated automatically per STU guidelines.", single_space=True, space_after=18)

    # LIST OF TABLES
    add_heading_1("LIST OF TABLES")
    lot_str = (
        "Table 3.1: System Hardware and Software Specification Requirements ........................ 28\n"
        "Table 3.2: Feasibility Analysis Matrix (Economic, Technical, Operational, Legal) .......... 32\n"
        "Table 3.3: Data Element Dictionary for Offline SQLite Database ............................ 44\n"
        "Table 3.4: Disease Classes, Scientific Names, and Sample Distributions ................... 47\n"
        "Table 3.5: Training Hyperparameter Configuration Summary .................................. 51\n"
        "Table 4.1: Per-Class Precision, Recall, and F1-Score Metrics on Held-Out Test Set .......... 61\n"
        "Table 4.2: Model Footprint, Latency, and Accuracy Benchmarks Across Runtimes .............. 63\n"
        "Table 4.3: Software System Integration Test Cases and Execution Outcomes ................... 68\n"
        "Table 4.4: System Usability Scale (SUS) Response Summary Across Participants .............. 71\n"
        "Table 4.5: State-of-the-Art Literature Comparison Matrix .................................. 74\n"
        "Table 4.6: Research Gap Fulfillment Verification Matrix ................................... 76"
    )
    add_p(lot_str, single_space=True, space_after=18)

    # LIST OF FIGURES
    add_heading_1("LIST OF FIGURES")
    lof_str = (
        "Figure 3.1: High-Level System Architecture and Offline Inference Pipeline .................. 34\n"
        "Figure 3.2: Context Data Flow Diagram (Level 0 DFD) ........................................ 37\n"
        "Figure 3.3: Level 1 Data Flow Diagram for Image Diagnostic & Remedy Retrieval ............... 38\n"
        "Figure 3.4: Use Case Diagram for Smallholder Farmers and Extension Officers .................. 40\n"
        "Figure 3.5: System Activity Flowchart for Camera Scan and Offline Treatment Matching ........ 41\n"
        "Figure 3.6: Entity-Relationship Diagram (ERD) for Diagnostic Logs and Remedies .............. 43\n"
        "Figure 3.7: Triplet Attention Module Architecture (Channel & Spatial Dimensions) .......... 49\n"
        "Figure 4.1: Model Training and Validation Accuracy/Loss Curves ............................ 60\n"
        "Figure 4.2: Test Set Classification Precision and Loss Metrics ............................. 61\n"
        "Figure 4.3: Grad-CAM Feature Map Attention Heatmaps Across Pathological Features .............. 65\n"
        "Figure 4.4: Flutter Mobile Client Home Screen & Camera Scanner UI .......................... 69\n"
        "Figure 4.5: Offline Diagnostic Report Page with Agronomic Treatment Remedies ................ 70\n"
        "Figure 4.6: Historical Batch Evaluation Logs and SQLite Database Records .................. 70"
    )
    add_p(lof_str, single_space=True, space_after=18)

    # =========================================================================
    # ------------------ CHAPTER ONE: INTRODUCTION ----------------------------
    # =========================================================================
    add_heading_1("CHAPTER ONE")
    add_heading_2("INTRODUCTION")

    add_heading_2("1.1 Background of the Study / Project")
    add_p(
        "Agriculture forms the cornerstone of Ghana's economy, employing over 40% of the active labor force and contributing "
        "significantly to the Gross Domestic Product (GDP). In the Bono Region, particularly within the Sunyani Municipality "
        "and surrounding farming enclaves such as Fiapre, Abesim, Chiraa, and Odumase, smallholder farmers cultivate staple food crops "
        "including tomato (Solanum lycopersicum) and maize (Zea mays). Tomato serves as an indispensable ingredient in Ghanaian cuisine, "
        "while maize represents the primary cereal crop supporting domestic food security and livestock feed."
    )
    add_p(
        "Despite their strategic importance, crop production in Ghana is perpetually threatened by plant diseases caused by fungal, "
        "bacterial, and viral pathogens. Infections such as Tomato Yellow Leaf Curl Virus, Early Blight, Late Blight, Septoria Leaf Spot, "
        "and Maize Northern Leaf Blight frequently cause devastating yield losses ranging from 30% to 100% if not detected and treated early. "
        "Traditionally, farmers rely on visual inspections by Ministry of Food and Agriculture (MoFA) agricultural extension officers. "
        "However, Ghana faces a severe deficit in extension personnel, with an estimated ratio of 1 extension officer to over 1,500 smallholder farmers. "
        "Consequently, farmers either misdiagnose crop ailments, leading to improper chemical application, or detect infections too late to salvage their harvests."
    )
    add_p(
        "Recent advances in Artificial Intelligence (AI), specifically Deep Convolutional Neural Networks (CNNs), have demonstrated remarkable success "
        "in automated image recognition and computer vision. However, deploying these deep learning models to solve real-world agricultural problems in "
        "rural Ghana presents significant technical hurdles. Most modern neural networks require high-performance GPU hardware or cloud servers, "
        "demanding continuous internet access. In rural Sunyani farming communities, cellular coverage is sparse or non-existent. Furthermore, large "
        "model sizes lead to memory bloat, high computational latency, and severe battery drain on low-cost Android smartphones common among farmers."
    )

    add_heading_2("1.2 Statement of Problem")
    add_p(
        "Smallholder farmers in the Sunyani Municipality suffer severe economic losses due to late and inaccurate crop disease diagnosis. "
        "Existing agricultural digital tools exhibit five fundamental deficiencies:\n"
        "1. Cloud Network Dependency: Most commercial AI diagnostic applications process images on remote cloud servers, rendering them unusable in internet-deprived rural farming zones.\n"
        "2. Excessive Model Size & Processing Delays: Standard deep neural networks (e.g., ResNet-50, DenseNet-121) exceed 100 MB, causing heavy memory consumption and sub-second delays on budget mobile processors.\n"
        "3. Lack of Visual Explainability: Deep learning classifiers operate as 'black boxes', outputting disease labels without explaining which leaf features influenced the prediction, breeding distrust among farmers and extension staff.\n"
        "4. Absence of Immediate Actionable Treatment Advice: Diagnostic apps frequently display disease names without providing localized, low-cost organic and chemical remedies.\n"
        "5. Fragmented Mobile Software Architecture: Machine learning research scripts rarely get converted into complete, production-ready, user-tested cross-platform mobile software applications."
    )

    add_heading_2("1.3 Objectives of the Study")
    add_heading_3("1.3.1 General Objective")
    add_p(
        "To design, develop, and empirically evaluate a hybrid offline-first mobile crop disease detection and advisory system "
        "leveraging a lightweight Triplet Attention EfficientNet-B0 architecture, INT8 post-training quantization, Grad-CAM visual explainability, "
        "and an embedded SQLite treatment database."
    )
    
    add_heading_3("1.3.2 Specific SMART Objectives")
    add_p("1. Dataset Curation & Preprocessing: To curate and balance a benchmark dataset of 21,394 tomato and maize leaf images across 14 pathological classes.")
    add_p("2. Neural Network Architecture Engineering: To design and train a lightweight EfficientNet-B0 backbone enhanced with a Triplet Attention Mechanism for spatial and cross-channel feature interaction.")
    add_p("3. Model Compression & Quantization: To apply post-training INT8 symmetric quantization to compress the model binary below 6 MB and achieve sub-100ms mobile CPU latency.")
    add_p("4. Visual Explainability (XAI): To implement Gradient-weighted Class Activation Mapping (Grad-CAM) for real-time visual lesion heatmap visualization.")
    add_p("5. Mobile Software Engineering: To build a cross-platform mobile application in Flutter backed by an embedded SQLite database providing offline chemical, organic, and cultural remedies.")
    add_p("6. Field Usability Evaluation: To evaluate system performance and user satisfaction with 15 target end-users in Sunyani using the System Usability Scale (SUS).")

    add_heading_3("1.3.3 Research Questions")
    add_p("RQ1: How much compression can post-training INT8 quantization achieve on a Triplet Attention EfficientNet-B0 model without causing significant loss in diagnostic accuracy?")
    add_p("RQ2: What is the on-device execution latency of the quantized TFLite model on budget mobile processors without internet connectivity?")
    add_p("RQ3: Can Grad-CAM heatmaps accurately localize pathological lesion features on leaf surfaces to establish visual trust?")
    add_p("RQ4: What is the practical usability rating of the offline mobile app among smallholder farmers and extension officers in Sunyani?")

    add_heading_2("1.4 Scope of the Project")
    add_p(
        "The geographical scope of field evaluation is centered within the Sunyani Municipality and Sunyani West District in the Bono Region of Ghana. "
        "The crop domain focuses specifically on tomato (Solanum lycopersicum) and maize (Zea mays), covering 14 distinct pathological classes "
        "(10 tomato categories and 4 maize categories). Software deployment targets native Android mobile devices running Android 7.0 (API level 24) "
        "and above, operating strictly offline without external web API dependencies."
    )

    add_heading_2("1.5 Significance of the Project")
    add_p(
        "From an agronomic perspective, this system equips smallholder farmers with an instant, expert-level diagnostic tool in their pockets, "
        "enabling timely disease intervention and reducing harvest losses. From a computer science and software engineering perspective, "
        "this work demonstrates a complete hybrid methodology—bridging advanced machine learning research (attention mechanisms, quantization, XAI) "
        "with robust software engineering practices (requirements specifications, DFDs, ERD, offline mobile architecture)."
    )

    add_heading_2("1.6 Limitations and Delimitations")
    add_heading_3("1.6.1 Limitations")
    add_p("1. Hardware Boundaries: Testing was limited to budget Android smartphones (Tecno, Infinix, Samsung) common in West Africa.")
    add_p("2. Single-Leaf Framing: The neural network requires close-up photos of individual leaves rather than full crop canopy shots.")
    add_p("3. Environmental Factors: Extreme outdoor glare and wet leaf surfaces can occasionally affect feature extraction quality.")

    add_heading_3("1.6.2 Delimitations")
    add_p("1. Excluded Crops: Major regional staple crops such as cassava, yam, plantain, and cocoa were excluded from the initial version.")
    add_p("2. Platform Target: iOS deployment was omitted during field trials to focus on Android devices predominant in rural Ghana.")

    add_heading_2("1.7 Organization of Work")
    add_p(
        "This project report is organized into five structured chapters: Chapter 1 introduces the background, problem statement, and objectives. "
        "Chapter 2 reviews related literature and highlights research gaps. Chapter 3 presents the hybrid methodology—covering Software System Analysis & Design (Part A) "
        "and Scientific AI Research Methodology (Part B). Chapter 4 presents system implementation, experimental results, testing outcomes, and discussion. "
        "Chapter 5 concludes the report with a summary of findings, limitations, and recommendations for future work."
    )

    # Save checkpoint
    doc.save(str(doc_path))
    print(f"Chapter 1 written. Total paragraphs: {len(doc.paragraphs)}")

if __name__ == '__main__':
    build_full_hybrid_thesis()
