import os
import sys
import shutil
from pathlib import Path

base_dir = Path(__file__).resolve().parent
sys_path_filtered = [p for p in sys.path if p != '' and Path(p).resolve() != base_dir]
sys.path = sys_path_filtered
import docx
from docx.shared import Inches, Pt

def enrich_chapter_5():
    doc_path = base_dir / "crop_disease_detection-stu-ch4-verified.docx"
    if not doc_path.exists():
        doc_path = base_dir / "crop_disease_detection-final-stu.docx"

    doc = docx.Document(str(doc_path))

    # Find Chapter Five paragraph
    ch5_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().upper()
        if txt in ["CHAPTER FIVE", "CHAPTER 5"]:
            ch5_idx = idx
            break

    if ch5_idx != -1:
        p_ch5_title = doc.paragraphs[ch5_idx]
        p_ch5_title.text = "CHAPTER FIVE"
        p_ch5_title.paragraph_format.space_before = Pt(18)
        p_ch5_title.paragraph_format.space_after = Pt(4)
        for r in p_ch5_title.runs:
            r.bold = True
            r.font.name = 'Times New Roman'

        # Ensure heading is SUMMARY, CONCLUSION AND RECOMMENDATIONS
        if ch5_idx + 1 < len(doc.paragraphs):
            p_ch5_sub = doc.paragraphs[ch5_idx + 1]
            p_ch5_sub.text = "SUMMARY, CONCLUSION AND RECOMMENDATIONS"
            p_ch5_sub.paragraph_format.space_before = Pt(4)
            p_ch5_sub.paragraph_format.space_after = Pt(12)
            for r in p_ch5_sub.runs:
                r.bold = True
                r.font.name = 'Times New Roman'

        # Insert STU Chapter 5 Intro Preamble
        stu_ch5_preamble = (
            "This final chapter presents a comprehensive synthesis of our research project. "
            "It is structured into three main sections to inform the reader of the project's conclusions and future prospects:\n"
            "• Section 5.1 (Summary): Provides a consolidated summary of the research background, problem statement, hybrid methodology, deep learning model engineering, quantization, mobile application development, and key experimental findings.\n"
            "• Section 5.2 (Conclusion): Details the principal conclusions derived directly from our empirical findings, structured into dedicated paragraphs corresponding to each specific research objective outlined in Chapter 1.\n"
            "• Section 5.3 (Recommendations): Offers strategic recommendations and technical directions for future research and expanded field deployment."
        )

        if ch5_idx + 2 < len(doc.paragraphs) and "This final chapter presents a comprehensive synthesis" not in doc.paragraphs[ch5_idx + 2].text:
            p_prem = doc.paragraphs[ch5_idx + 1].insert_paragraph_before(stu_ch5_preamble, style='Normal')
            p_prem.paragraph_format.space_before = Pt(6)
            p_prem.paragraph_format.space_after = Pt(12)
            p_prem.paragraph_format.line_spacing = 1.5

    # Re-build Sections 5.1, 5.2, and 5.3 to strictly follow STU guidelines
    # Find paragraph indices for 5.1, 5.2, 5.3
    s51_idx = -1
    s52_idx = -1
    s53_idx = -1
    app_idx = -1

    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if "5.1 Summary" in txt:
            s51_idx = idx
        elif "5.2 Conclusion" in txt:
            s52_idx = idx
        elif "5.3 Recommendation" in txt:
            s53_idx = idx
        elif "APPENDIX" in txt and app_idx == -1:
            app_idx = idx

    # If 5.1 Summary exists, replace its body text with full STU summary
    if s51_idx != -1:
        p_51 = doc.paragraphs[s51_idx]
        p_51.text = "5.1 Summary"
        p_51.paragraph_format.space_before = Pt(14)
        p_51.paragraph_format.space_after = Pt(4)
        for r in p_51.runs:
            r.bold = True
            r.font.name = 'Times New Roman'

        summary_body = (
            "This research project addressed a critical agricultural challenge in Ghana: the severe economic crop losses suffered by smallholder farmers in the Sunyani Municipality due to late and inaccurate plant disease diagnosis. "
            "Smallholder farmers in rural communities face two major structural bottlenecks: an acute shortage of agricultural extension officers (ratio of 1:1,500+) and non-existent cellular network coverage in remote farming enclaves. "
            "To solve these challenges, our project team designed, built, and evaluated an offline-first mobile crop disease detection and advisory system focused on tomato (Solanum lycopersicum) and maize (Zea mays).\n\n"
            "Our technical methodology combined advanced deep learning research with robust mobile software engineering. We integrated a Triplet Attention Mechanism into a lightweight EfficientNet-B0 backbone to extract spatial and cross-channel lesion features. "
            "We trained and validated the neural network on a curated dataset of 21,394 leaf images spanning 14 pathological categories (10 tomato classes and 4 maize classes) using a two-stage transfer learning schedule. "
            "To enable seamless offline execution on budget Android smartphones, we applied post-training INT8 symmetric quantization, compressing the model binary from 20.3 MB down to 5.1 MB (a 74.8% memory savings) while keeping test accuracy high at 97.85% (compared to the 98.24% FP32 baseline).\n\n"
            "Operating inside a cross-platform Flutter application backed by an embedded TensorFlow Lite runtime, on-device inference latency averaged 92 ms per photo without requiring any cellular internet connection. "
            "We integrated Gradient-weighted Class Activation Mapping (Grad-CAM) to output visual attention heatmaps, providing transparent explainability. Furthermore, an embedded local SQLite database supplies immediate organic, chemical, and cultural treatment advice. "
            "Empirical usability testing with 15 target users in Sunyani (10 smallholder farmers and 5 extension officers) yielded a System Usability Scale (SUS) score of 76.5 out of 100 (Grade B, 'Good'), demonstrating strong practical utility."
        )
        
        # Replace text of paragraph after 5.1
        if s51_idx + 1 < len(doc.paragraphs):
            doc.paragraphs[s51_idx + 1].text = summary_body
            doc.paragraphs[s51_idx + 1].paragraph_format.space_after = Pt(12)
            doc.paragraphs[s51_idx + 1].paragraph_format.line_spacing = 1.5

    # If 5.2 Conclusion exists, format paragraph by paragraph for EACH specific objective
    if s52_idx != -1:
        p_52 = doc.paragraphs[s52_idx]
        p_52.text = "5.2 Conclusion"
        p_52.paragraph_format.space_before = Pt(14)
        p_52.paragraph_format.space_after = Pt(4)
        for r in p_52.runs:
            r.bold = True
            r.font.name = 'Times New Roman'

        conclusion_intro = "Based on the empirical findings and experimental evaluation of this study, our principal conclusions are structured directly corresponding to each of the six specific research objectives outlined in Chapter 1:"

        # Paragraph for Objective 1
        conc_obj1 = (
            "Conclusion on Objective 1 (Dataset Curation and Stratification): "
            "Curating a balanced dataset of 21,394 leaf images across 14 pathological classes and employing a 70/15/15 stratified split provided a robust foundation for neural network training. "
            "Isolating 3,209 held-out test images ensured high statistical reliability, confirming that rigorous data stratification is essential for training deep learning models that generalize well across varied agricultural field conditions."
        )

        # Paragraph for Objective 2
        conc_obj2 = (
            "Conclusion on Objective 2 (Neural Network Engineering and Triplet Attention): "
            "Integrating a Triplet Attention Mechanism into an EfficientNet-B0 backbone significantly enhanced spatial and cross-channel lesion feature representation. "
            "Achieving a test accuracy of 98.24% on the FP32 baseline proves that capturing cross-dimension feature dependencies allows lightweight networks to achieve diagnostic precision comparable to massive deep architectures."
        )

        # Paragraph for Objective 3
        conc_obj3 = (
            "Conclusion on Objective 3 (Model Compression and INT8 Quantization): "
            "Applying post-training INT8 symmetric quantization successfully compressed the model binary from 20.3 MB down to 5.1 MB—a 74.8% memory reduction—while incurring only a slight 0.39% accuracy trade-off (97.85% INT8 vs. 98.24% FP32). "
            "Achieving an average on-device CPU execution latency of 92 ms on budget smartphones proves that 8-bit quantization enables complex neural networks to operate locally on low-cost hardware without cloud server dependencies."
        )

        # Paragraph for Objective 4
        conc_obj4 = (
            "Conclusion on Objective 4 (Visual Explainability via Grad-CAM Heatmaps): "
            "Implementing Gradient-weighted Class Activation Mapping (Grad-CAM) successfully resolved the 'black-box' opacity of deep learning models. "
            "Generating real-time visual attention heatmaps that highlight key pathological features (such as Septoria spots and Yellow Leaf Curl chlorosis) establishes visual trust among farmers and agricultural extension officers."
        )

        # Paragraph for Objective 5
        conc_obj5 = (
            "Conclusion on Objective 5 (Mobile Software Engineering and Offline SQLite Database): "
            "Decoupling the machine learning inference module from local SQLite treatment storage created a highly resilient, offline-first mobile software application. "
            "Supplying immediate, localized chemical, organic, and cultural remedies directly on the phone solves the critical information gap faced by rural farmers operating in areas without cellular connectivity."
        )

        # Paragraph for Objective 6
        conc_obj6 = (
            "Conclusion on Objective 6 (Field Usability Evaluation): "
            "Conducting empirical field evaluation with 15 target users in Sunyani yielded a System Usability Scale (SUS) score of 76.5 out of 100 (Grade B, 'Good'). "
            "This confirms that the application possesses strong practical usability, intuitive navigation, and high field acceptability among smallholder farmers and extension personnel."
        )

        # Replace text in paragraphs under 5.2
        p_intro = doc.paragraphs[s52_idx + 1]
        p_intro.text = conclusion_intro + "\n\n" + conc_obj1 + "\n\n" + conc_obj2 + "\n\n" + conc_obj3 + "\n\n" + conc_obj4 + "\n\n" + conc_obj5 + "\n\n" + conc_obj6
        p_intro.paragraph_format.space_after = Pt(12)
        p_intro.paragraph_format.line_spacing = 1.5

    # If 5.3 Recommendation(s) exists, format future work recommendations
    if s53_idx != -1:
        p_53 = doc.paragraphs[s53_idx]
        p_53.text = "5.3 Recommendation(s)"
        p_53.paragraph_format.space_before = Pt(14)
        p_53.paragraph_format.space_after = Pt(4)
        for r in p_53.runs:
            r.bold = True
            r.font.name = 'Times New Roman'

        recs_text = (
            "Based on the outcomes and boundaries of this research, we offer the following recommendations for future academic research and field engineering:\n\n"
            "1. Expansion to Additional Regional Staple Crops: Future work should expand the dataset and offline treatment dictionary to include major West African staple crops such as cassava (African cassava mosaic virus, brown streak), cocoa (black pod disease, swollen shoot virus), yam, and plantain.\n\n"
            "2. Implementation of Real-Time Video Object Detection: Future technical iterations should integrate lightweight object detection architectures (such as YOLOv8-nano TFLite) to enable real-time lesion scanning and bounding-box detection within live mobile camera video streams.\n\n"
            "3. Integration of Local Ghanaian Voice Advisory Support: To assist low-literacy farmers in rural communities, future software updates should integrate offline Text-to-Speech (TTS) audio narration supplying treatment remedies in major local Ghanaian languages, including Twi, Fante, Ewe, and Dagbani.\n\n"
            "4. Multi-Modal Environmental Sensor Data Fusion: Future research could explore combining image-based leaf diagnostics with low-cost ambient microclimate sensors (temperature, humidity, soil moisture) to predict disease outbreak risks before visual symptoms appear."
        )
        
        if s53_idx + 1 < len(doc.paragraphs):
            doc.paragraphs[s53_idx + 1].text = recs_text
            doc.paragraphs[s53_idx + 1].paragraph_format.space_after = Pt(12)
            doc.paragraphs[s53_idx + 1].paragraph_format.line_spacing = 1.5

    # Save to crop_disease_detection-stu-ch5-verified.docx
    out_file = base_dir / "crop_disease_detection-stu-ch5-verified.docx"
    doc.save(str(out_file))
    print(f"Successfully saved verified Chapter 5 document to: {out_file}")

    for fname in ["crop_disease_detection-final-stu.docx", "crop_disease_detection-real-perfect.docx", "crop_disease_detection-real.docx"]:
        try:
            doc.save(str(base_dir / fname))
            print(f"Also saved copy to: {fname}")
        except Exception:
            pass

    print("\n==================================================")
    print("      CHAPTER 5 STU VERIFICATION SUMMARY          ")
    print("==================================================")
    print(f"Final Paragraph Count: {len(doc.paragraphs)}")
    print(f"Final Table Count: {len(doc.tables)}")
    total_words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"Final Total Word Count: {total_words}")
    print(f"Estimated Double-Spaced Pages: {round(total_words / 250, 1)}")
    print("==================================================")

if __name__ == '__main__':
    enrich_chapter_5()
