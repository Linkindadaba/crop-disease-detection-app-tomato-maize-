import os
import sys
import shutil
from pathlib import Path

base_dir = Path(__file__).resolve().parent
sys_path_filtered = [p for p in sys.path if p != '' and Path(p).resolve() != base_dir]
sys.path = sys_path_filtered
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_document():
    # Load backup if available or original
    doc_path = base_dir / "crop_disease_detection-real-backup.docx"
    if not doc_path.exists():
        doc_path = base_dir / "crop_disease_detection-real.docx"

    doc = docx.Document(str(doc_path))
    body = doc.paragraphs[0]._element.getparent()
    
    print(f"Initial body children count: {len(body)}")
    print(f"Initial paragraphs count: {len(doc.paragraphs)}")
    print(f"Initial tables count: {len(doc.tables)}")

    # 1. Standardize Margins (1.5" Left for STU binding, 1.0" Top, Bottom, Right)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)

    # 2. Identify & Remove Corrupted/Extraneous Tables in Chapter 3
    tables_to_remove = []
    for idx, tbl in enumerate(doc.tables):
        cols = len(tbl.columns)
        rows = len(tbl.rows)
        hdr_text = " ".join([c.text.strip() for c in tbl.rows[0].cells[:4]])
        
        # Corrupted tables #7, #8, #9 have 12 columns and repeated headers
        if cols >= 10 or rows > 25:
            if "Precision" in hdr_text and ("Disease Class" in hdr_text or "Disease Category" in hdr_text) and rows in [134, 136, 29, 31]:
                tables_to_remove.append(tbl)
                print(f"Flagged corrupted table #{idx+1} ({rows} rows x {cols} cols) for removal.")

    for tbl in tables_to_remove:
        elem = tbl._element
        elem.getparent().remove(elem)
    print(f"Removed {len(tables_to_remove)} corrupted tables from Chapter 3.")

    # 3. Remove Empty Paragraph Gaps (Blank line paragraphs)
    empty_paragraphs_removed = 0
    paragraphs_to_remove = []
    
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # Check if paragraph has images/pictures attached
        has_drawing = 'w:drawing' in p._element.xml or 'w:pict' in p._element.xml
        
        if not text and not has_drawing:
            paragraphs_to_remove.append(p)
            empty_paragraphs_removed += 1

    for p in paragraphs_to_remove:
        elem = p._element
        if elem.getparent() is not None:
            elem.getparent().remove(elem)

    print(f"Removed {empty_paragraphs_removed} empty paragraph gaps from document.")

    # 4. Clean Paragraph Spacing and Formatting
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
            
        txt_upper = txt.upper()
        
        # Chapter headings
        if txt_upper in ["CHAPTER ONE", "CHAPTER TWO", "CHAPTER THREE", "CHAPTER FOUR", "CHAPTER FIVE", "REFERENCES", "DECLARATION", "CERTIFICATION", "DEDICATION", "ACKNOWLEDGEMENT", "ABSTRACT", "TABLE OF CONTENTS", "LIST OF TABLES", "LIST OF FIGURES"]:
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.5
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
        # Section headings
        elif any(p.style.name.startswith(h) for h in ['Heading 1', 'Heading 2', 'Heading 3']) or (len(txt) > 3 and txt[0].isdigit() and '.' in txt[:4]):
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.5
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
        # Normal body text
        else:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5
            for r in p.runs:
                r.font.name = 'Times New Roman'

    # Save cleaned document to both crop_disease_detection-real-perfect.docx and attempt original
    out_perfect = base_dir / "crop_disease_detection-real-perfect.docx"
    doc.save(str(out_perfect))
    print(f"Saved pristine cleaned document to: {out_perfect}")

    out_orig = base_dir / "crop_disease_detection-real.docx"
    try:
        doc.save(str(out_orig))
        print(f"Successfully updated original file: {out_orig}")
    except Exception as e:
        print(f"Could not overwrite {out_orig} directly (File might be open in MS Word). User can view {out_perfect}.")

    print("\n==================================================")
    print("      CLEANED DOCUMENT SUMMARY VERIFICATION      ")
    print("==================================================")
    print(f"Final Body Children Count: {len(body)}")
    print(f"Final Paragraph Count: {len(doc.paragraphs)}")
    print(f"Final Table Count: {len(doc.tables)}")
    total_words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"Final Total Word Count: {total_words}")
    print(f"Estimated Double-Spaced Pages: {round(total_words / 250, 1)}")
    print("==================================================")

if __name__ == '__main__':
    clean_document()
