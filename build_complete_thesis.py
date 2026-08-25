import os
import sys
import shutil
from pathlib import Path

base_dir = Path(__file__).resolve().parent

# Ensure images exist in media folder
screenshots_dir = base_dir / "app_screenshots"
media_dir = base_dir / "media"
media_dir.mkdir(exist_ok=True)

img_map = {
    "Screenshot 2026-08-01 160612.png": "image12.png",
    "Diagnose_report_Screenshot.png": "image13.png",
    "full_report_Screenshot.png": "image14.png"
}

for src_name, dst_name in img_map.items():
    src_file = screenshots_dir / src_name
    dst_file = media_dir / dst_name
    if src_file.exists():
        shutil.copy(src_file, dst_file)

# Filter sys.path to avoid local docx folder shadowing python-docx package
sys_path_filtered = [p for p in sys.path if p != '' and Path(p).resolve() != base_dir]
sys.path = sys_path_filtered
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    # 1. First ensure expand_doc_to_50_pages.py runs
    expand_script = base_dir / "expand_doc_to_50_pages.py"
    if expand_script.exists():
        import subprocess
        subprocess.run([sys.executable, str(expand_script)], check=False)

    # 2. Next ensure update_stu_compliance.py runs
    compliance_script = base_dir / "update_stu_compliance.py"
    if compliance_script.exists():
        import subprocess
        subprocess.run([sys.executable, str(compliance_script)], check=False)

    # 3. Read crop_disease_detection.docx and verify contents
    doc_path = base_dir / "crop_disease_detection.docx"
    doc = docx.Document(str(doc_path))
    
    total_words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"==================================================")
    print(f"FINAL HYBRID THESIS GENERATION VERIFICATION")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Words: {total_words}")
    print(f"Estimated Double-Spaced Pages: {round(total_words / 250, 1)}")
    print(f"==================================================")

if __name__ == '__main__':
    main()
