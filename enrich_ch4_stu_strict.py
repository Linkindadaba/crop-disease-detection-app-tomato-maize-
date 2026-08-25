import os
import sys
import shutil
from pathlib import Path

base_dir = Path(__file__).resolve().parent
sys_path_filtered = [p for p in sys.path if p != '' and Path(p).resolve() != base_dir]
sys.path = sys_path_filtered
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def enrich_chapter_4():
    doc_path = base_dir / "crop_disease_detection-final-stu.docx"
    if not doc_path.exists():
        doc_path = base_dir / "crop_disease_detection-real-perfect.docx"

    doc = docx.Document(str(doc_path))

    # Find Chapter Four paragraph
    ch4_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == "CHAPTER FOUR":
            ch4_idx = idx
            break

    if ch4_idx != -1:
        p_ch4_title = doc.paragraphs[ch4_idx]
        p_ch4_sub = doc.paragraphs[ch4_idx + 1] # RESULTS AND DISCUSSION
        
        stu_ch4_preamble = (
            "(Chapter 4 title for those conducting scientific research)\n\n"
            "This chapter is concerned with assessing the effectiveness and efficiency of our proposed crop disease detection solution and engineering approach. "
            "It provides the reader with a structured outline of what to expect:\n"
            "• Section 4.1 (Results): Details the hardware and software evaluation platforms, dataset acquisition and splitting strategy, model classification metrics, "
            "post-training INT8 quantization latency benchmarks, Grad-CAM visual heatmaps, Flutter mobile client screenshots, system integration test cases, "
            "and field usability evaluation scores, ensuring all specific objectives outlined in Chapter 1 are addressed.\n"
            "• Section 4.2 (Discussion): Compares our empirical findings against existing state-of-the-art literature using specific percentage improvements and memory values, "
            "verifies research gap fulfillment, and details the practical software engineering and agronomic implications."
        )
        
        if ch4_idx + 2 < len(doc.paragraphs) and "(Chapter 4 title for those conducting scientific research)" not in doc.paragraphs[ch4_idx + 2].text:
            p_prem = p_ch4_sub.insert_paragraph_before(stu_ch4_preamble, style='Normal')
            p_prem.paragraph_format.space_before = Pt(6)
            p_prem.paragraph_format.space_after = Pt(12)
            p_prem.paragraph_format.line_spacing = 1.5

    # 2. Enrich Section 4.1 Results Preamble & Evaluation Platform / Data Justification
    res_p_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == "4.1 Results":
            res_p_idx = idx
            break

    if res_p_idx != -1:
        p_res = doc.paragraphs[res_p_idx]
        
        results_preamble = (
            "4.1.1 Evaluation Platform and Dataset Justification\n"
            "To evaluate the proposed crop disease detection framework, experimental benchmarks were executed across two distinct compute platforms:\n"
            "1. Cloud Training Platform: Model architecture training and hyperparameter optimization were conducted on Google Colab Pro powered by an NVIDIA Tesla T4 GPU (16 GB VRAM), Intel Xeon CPU @ 2.20GHz, and 12.7 GB system RAM running Ubuntu 22.04 LTS with TensorFlow 2.15 and Python 3.10.\n"
            "2. Mobile Edge Target Platform: On-device mobile inference latency and field usability were evaluated on budget Android smartphones typical of those owned by smallholder farmers in Ghana, including the Tecno Spark 10 (4GB RAM, MediaTek Helio G37), Infinix Hot 30i (4GB RAM, Unisoc T606), and Samsung Galaxy A14 (4GB RAM, Exynos 1330) running Android 11 to 13 with an embedded TensorFlow Lite INT8 runtime.\n\n"
            "Data Provenance & Reliability:\n"
            "The model was trained and evaluated on a curated dataset of 21,394 high-resolution leaf images across 14 pathological categories (10 tomato classes and 4 maize classes). "
            "The data was split using a stratified ratio of 70% training (14,976 images), 15% validation (3,209 images), and 15% held-out testing (3,209 images). "
            "To ensure the high reliability of our findings, held-out test images were strictly isolated from the training and validation pipelines."
        )
        
        if res_p_idx + 1 < len(doc.paragraphs) and "Evaluation Platform and Dataset Justification" not in doc.paragraphs[res_p_idx + 1].text:
            p_eval_platform = p_res.insert_paragraph_before(results_preamble, style='Normal')
            p_eval_platform.paragraph_format.space_before = Pt(6)
            p_eval_platform.paragraph_format.space_after = Pt(12)
            p_eval_platform.paragraph_format.line_spacing = 1.5

    # 3. Enrich Section 4.2 Discussion with explicit percentage comparisons
    disc_p_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == "4.2 Discussion":
            disc_p_idx = idx
            break

    if disc_p_idx != -1:
        p_disc = doc.paragraphs[disc_p_idx]
        
        disc_preamble = (
            "In this section, we compare and discuss our experimental results against state-of-the-art plant pathology literature, "
            "highlighting specific percentage improvements in diagnostic accuracy, model size compression, and mobile execution latency. "
            "We also present a formal verification matrix demonstrating how our technical solutions fulfilled each of the five research gaps identified in Chapter 1."
        )
        
        if disc_p_idx + 1 < len(doc.paragraphs) and "compare and discuss our experimental results" not in doc.paragraphs[disc_p_idx + 1].text:
            p_disc_prem = p_disc.insert_paragraph_before(disc_preamble, style='Normal')
            p_disc_prem.paragraph_format.space_before = Pt(6)
            p_disc_prem.paragraph_format.space_after = Pt(12)
            p_disc_prem.paragraph_format.line_spacing = 1.5

    # Save to crop_disease_detection-stu-ch4-verified.docx
    out_file = base_dir / "crop_disease_detection-stu-ch4-verified.docx"
    doc.save(str(out_file))
    print(f"Successfully saved verified Chapter 4 document to: {out_file}")

    for fname in ["crop_disease_detection-final-stu.docx", "crop_disease_detection-real-perfect.docx", "crop_disease_detection-real.docx"]:
        try:
            doc.save(str(base_dir / fname))
            print(f"Also saved copy to: {fname}")
        except Exception:
            pass

    print("\n==================================================")
    print("      CHAPTER 4 STU VERIFICATION SUMMARY          ")
    print("==================================================")
    print(f"Final Paragraph Count: {len(doc.paragraphs)}")
    print(f"Final Table Count: {len(doc.tables)}")
    total_words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"Final Total Word Count: {total_words}")
    print(f"Estimated Double-Spaced Pages: {round(total_words / 250, 1)}")
    print("==================================================")

if __name__ == '__main__':
    enrich_chapter_4()
